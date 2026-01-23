#!/usr/bin/env python3
"""
Convert a CAT paper metadata DOCX into the JSON schema used by scripts/import-paper.mjs
and optionally upload it to Supabase.

Usage:
  python scripts/convert_paper_metadata.py --docx "/path/to/PAPER METADATA_CAT2024_slot1.docx" --out data/cat-2024-slot1.json --publish

Dependencies:
  pip install python-docx supabase python-dotenv

Env vars (same as import-paper.mjs):
  NEXT_PUBLIC_SUPABASE_URL
  SUPABASE_SERVICE_ROLE_KEY
"""

import argparse
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

from docx import Document
from supabase import create_client, Client

# -----------------------------------------------------------------------------
# Parsing helpers
# -----------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _cell_text(cell) -> str:
    return _clean_text("\n".join(p.text for p in cell.paragraphs))


def _normalize_key(key: str) -> str:
    key = re.sub(r"[^A-Za-z0-9]+", "_", key).strip("_").lower()
    return key


def parse_paper_metadata(table) -> Dict[str, Any]:
    meta: Dict[str, Any] = {}
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        key = _normalize_key(_cell_text(row.cells[0]))
        value = _cell_text(row.cells[1])
        if not key:
            continue
        meta[key] = value
    return meta


def parse_sections(meta: Dict[str, Any]) -> List[Dict[str, Any]]:
    sections: List[Dict[str, Any]] = []
    raw = meta.get("sections") or meta.get("section_breakup") or meta.get("sectionwise_breakup")
    if not raw:
        return sections

    lines = [p for chunk in raw.split("\n") for p in chunk.split(";") if p.strip()]
    for line in lines:
        parts = [p.strip() for p in re.split(r"[|,:]", line) if p.strip()]
        if len(parts) < 2:
            continue
        name = parts[0].upper()
        as_int = lambda x: int(float(x)) if x not in ("", None) else None  # noqa: E731
        try:
            questions = as_int(parts[1])
            time_val = as_int(parts[2]) if len(parts) > 2 else None
            marks_val = as_int(parts[3]) if len(parts) > 3 else None
        except ValueError:
            continue
        sections.append({
            "name": name,
            "questions": questions,
            "time": time_val,
            "marks": marks_val,
        })
    return sections


def detect_question_table(table) -> bool:
    headers = [_normalize_key(_cell_text(cell)) for cell in table.rows[0].cells]
    needed = {"question", "question_text", "question_number"}
    has_needed = any(h in needed for h in headers)
    has_answer = any("answer" in h for h in headers)
    return has_needed and has_answer


def parse_questions(tables: Sequence[Any]) -> List[Dict[str, Any]]:
    questions: List[Dict[str, Any]] = []
    for table in tables:
        if not table.rows:
            continue
        headers = [_normalize_key(_cell_text(cell)) for cell in table.rows[0].cells]
        if not detect_question_table(table):
            continue
        for row in table.rows[1:]:
            cells = row.cells
            row_map = {headers[i]: _cell_text(cells[i]) for i in range(min(len(headers), len(cells)))}
            q_text = row_map.get("question_text") or row_map.get("question")
            if not q_text:
                continue
            section = (row_map.get("section") or "").upper() or "VARC"
            q_type = (row_map.get("type") or row_map.get("question_type") or "MCQ").upper()
            q_num_raw = row_map.get("question_number") or row_map.get("qno") or row_map.get("sno")
            try:
                q_number = int(float(q_num_raw)) if q_num_raw else len(questions) + 1
            except ValueError:
                q_number = len(questions) + 1

            option_keys = [k for k in headers if re.match(r"option[a-d]", k)]
            options = [row_map.get(k) for k in option_keys if row_map.get(k)]
            if q_type == "MCQ" and not options:
                options = [row_map.get(k) for k in ["a", "b", "c", "d"] if row_map.get(k)]

            answer = row_map.get("answer") or row_map.get("correct_answer")
            difficulty = row_map.get("difficulty") or None
            topic = row_map.get("topic") or None
            subtopic = row_map.get("subtopic") or None
            solution = row_map.get("solution") or row_map.get("solution_text") or None

            q_entry = {
                "section": section or "VARC",
                "question_number": q_number,
                "question_text": q_text,
                "question_type": "TITA" if q_type == "TITA" else "MCQ",
                "options": options or None,
                "correct_answer": answer or "",
                "positive_marks": 3.0,
                "negative_marks": 0 if q_type == "TITA" else 1.0,
                "difficulty": difficulty,
                "topic": topic,
                "subtopic": subtopic,
                "solution_text": solution,
            }
            questions.append(q_entry)
    return questions


# -----------------------------------------------------------------------------
# Supabase upload
# -----------------------------------------------------------------------------

def get_supabase_client() -> Client:
    url = os.getenv("NEXT_PUBLIC_SUPABASE_URL") or os.getenv("SUPABASE_URL")
    key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
    if not url or not key:
        raise RuntimeError("Missing NEXT_PUBLIC_SUPABASE_URL or SUPABASE_SERVICE_ROLE_KEY")
    return create_client(url, key)


def upsert_paper(client: Client, paper: Dict[str, Any]) -> Dict[str, Any]:
    existing = client.table("papers").select("id").eq("slug", paper["slug"]).maybe_single().execute()
    rows = existing.data if hasattr(existing, "data") else existing.get("data")
    paper_payload = {
        "slug": paper["slug"],
        "title": paper["title"],
        "description": paper.get("description"),
        "year": int(paper.get("year") or datetime.now().year),
        "total_questions": int(paper.get("total_questions") or len(paper.get("questions", []))),
        "total_marks": paper.get("total_marks") or 198,
        "duration_minutes": paper.get("duration_minutes") or 120,
        "sections": paper.get("sections") or [],
        "default_positive_marks": float(paper.get("default_positive_marks") or 3.0),
        "default_negative_marks": float(paper.get("default_negative_marks") or 1.0),
        "difficulty_level": paper.get("difficulty_level") or "cat-level",
        "is_free": bool(paper.get("is_free", True)),
        "published": bool(paper.get("published", False)),
        "available_from": paper.get("available_from") or None,
        "available_until": paper.get("available_until") or None,
    }

    if rows:
        paper_id = rows[0]["id"]
        res = client.table("papers").update(paper_payload).eq("id", paper_id).execute()
        return {"id": paper_id, **paper_payload, "_supabase": res}
    res = client.table("papers").insert(paper_payload).execute()
    paper_id = res.data[0]["id"]
    return {"id": paper_id, **paper_payload, "_supabase": res}


def replace_questions(client: Client, paper_id: int, questions: List[Dict[str, Any]]) -> int:
    client.table("questions").delete().eq("paper_id", paper_id).execute()
    batch_size = 50
    total = 0
    for i in range(0, len(questions), batch_size):
        batch = questions[i:i + batch_size]
        payload = []
        for q in batch:
            payload.append({
                "paper_id": paper_id,
                "section": q.get("section"),
                "question_number": int(q.get("question_number", total + 1)),
                "question_text": q.get("question_text"),
                "question_type": q.get("question_type", "MCQ"),
                "options": q.get("options"),
                "correct_answer": q.get("correct_answer"),
                "positive_marks": q.get("positive_marks", 3.0),
                "negative_marks": q.get("negative_marks", 1.0),
                "difficulty": q.get("difficulty"),
                "topic": q.get("topic"),
                "subtopic": q.get("subtopic"),
                "solution_text": q.get("solution_text"),
                "solution_image_url": q.get("solution_image_url"),
                "video_solution_url": q.get("video_solution_url"),
                "is_active": True,
            })
        client.table("questions").insert(payload).execute()
        total += len(batch)
    return total


def publish_paper(client: Client, paper_id: int) -> None:
    client.table("papers").update({"published": True}).eq("id", paper_id).execute()


# -----------------------------------------------------------------------------
# Orchestration
# -----------------------------------------------------------------------------

def build_json(meta: Dict[str, Any], questions: List[Dict[str, Any]]) -> Dict[str, Any]:
    slug = meta.get("slug") or meta.get("paper_slug") or "cat-2024"
    title = meta.get("title") or meta.get("paper_title") or "CAT 2024"
    description = meta.get("description") or meta.get("desc")
    year_val = meta.get("year") or meta.get("exam_year") or datetime.now().year
    total_questions = meta.get("total_questions") or len(questions)
    total_marks = meta.get("total_marks") or None
    duration = meta.get("duration_minutes") or meta.get("duration") or 120

    sections = parse_sections(meta)
    if not sections and questions:
        counts: Dict[str, int] = {}
        for q in questions:
            sec = q.get("section", "VARC")
            counts[sec] = counts.get(sec, 0) + 1
        sections = [{"name": k, "questions": v, "time": None, "marks": None} for k, v in counts.items()]

    paper = {
        "slug": slug,
        "title": title,
        "description": description,
        "year": int(float(year_val)) if year_val else datetime.now().year,
        "total_questions": int(float(total_questions)) if total_questions else len(questions),
        "total_marks": int(float(total_marks)) if total_marks else None,
        "duration_minutes": int(float(duration)) if duration else 120,
        "sections": sections,
        "default_positive_marks": 3.0,
        "default_negative_marks": 1.0,
        "difficulty_level": meta.get("difficulty_level") or "cat-level",
        "is_free": True,
        "published": False,
        "available_from": None,
        "available_until": None,
    }

    return {"paper": paper, "questions": questions}


def convert_docx(docx_path: Path) -> Dict[str, Any]:
    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError("DOCX contains no tables to parse")

    paper_meta = parse_paper_metadata(doc.tables[0])
    question_tables = doc.tables[1:] if len(doc.tables) > 1 else []
    questions = parse_questions(question_tables)
    return build_json(paper_meta, questions)


def save_json(data: Dict[str, Any], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert CAT paper DOCX to JSON and upload to Supabase")
    parser.add_argument("--docx", required=True, help="Path to PAPER METADATA docx")
    parser.add_argument("--out", default="data/cat-2024.json", help="Where to write the JSON")
    parser.add_argument("--publish", action="store_true", help="Publish paper after upload")
    parser.add_argument("--upload", action="store_true", help="Upload to Supabase")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    out_path = Path(args.out)

    data = convert_docx(docx_path)
    save_json(data, out_path)
    print(f"✅ JSON written to {out_path}")

    if args.upload:
        client = get_supabase_client()
        paper = data["paper"]
        questions = data["questions"]
        supa_paper = upsert_paper(client, paper)
        count = replace_questions(client, supa_paper["id"], questions)
        if args.publish:
            publish_paper(client, supa_paper["id"])
        print(f"✅ Uploaded paper id={supa_paper['id']} with {count} questions; published={args.publish}")
    else:
        print("ℹ️ Skipped upload (run with --upload to push to Supabase)")


if __name__ == "__main__":
    main()
